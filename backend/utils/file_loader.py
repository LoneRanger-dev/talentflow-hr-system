import os
import io
import csv
from pathlib import Path
from typing import Dict, List, Any, Optional
from backend.config import JOB_DESC_FILE, RESUMES_DIR

def load_job_description() -> str:
    """Load job description markdown file."""
    if JOB_DESC_FILE.exists():
        with open(JOB_DESC_FILE, "r", encoding="utf-8") as f:
            return f.read()
    return "Job Description not found."

def save_job_description(content: str) -> bool:
    """Save/update job description content."""
    try:
        JOB_DESC_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(JOB_DESC_FILE, "w", encoding="utf-8") as f:
            f.write(content)
        return True
    except Exception as e:
        print(f"Error writing job description: {e}")
        return False

def extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from uploaded PDF, DOCX, TXT, CSV, or Markdown files."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    # 1. PDF File Parsing
    if ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    pages_text.append(extracted)
            if pages_text:
                return "\n\n".join(pages_text)
        except Exception as e:
            print(f"⚠️ PyPDF extraction error for {filename}: {e}")

    # 2. DOCX Word Document Parsing
    if ext in ["docx", "doc"]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            if paragraphs:
                return "\n".join(paragraphs)
        except Exception as e:
            print(f"⚠️ Docx extraction error for {filename}: {e}")

    # 3. CSV File Parsing
    if ext == "csv":
        try:
            decoded = file_bytes.decode("utf-8", errors="ignore")
            reader = csv.reader(io.StringIO(decoded))
            rows = [" ".join(row) for row in reader if row]
            return "\n".join(rows)
        except Exception as e:
            print(f"⚠️ CSV extraction error for {filename}: {e}")

    # 4. Standard Text / Markdown / Fallback
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"⚠️ Text decoding error for {filename}: {e}")
        return f"Resume text extracted from {filename}"

def load_all_resumes() -> Dict[str, Dict[str, Any]]:
    """Load all sample resume files from data/resumes directory."""
    resumes = {}
    if RESUMES_DIR.exists():
        resume_files = list(RESUMES_DIR.glob("*.*"))
        for file in resume_files:
            if file.suffix.lower() in [".md", ".markdown", ".txt", ".docx", ".pdf", ".csv"]:
                candidate_id = file.stem.replace("Resume - ", "").lower().replace(" ", "_")
                try:
                    with open(file, "rb") as f:
                        file_bytes = f.read()
                    content = extract_text_from_file_bytes(file_bytes, file.name)
                    resumes[candidate_id] = {
                        "filename": file.name,
                        "candidate_id": candidate_id,
                        "content": content,
                        "word_count": len(content.split())
                    }
                except Exception as e:
                    print(f"Error loading resume {file.name}: {e}")
    return resumes
